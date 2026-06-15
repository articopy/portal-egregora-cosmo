import os
import docx

TEMPLATE_PATH = r"c:\Projetos\Cosmo Alma TV\portal-egregora\Contrato\CONTRATO_COSMO_ALMA_TV_V3.docx"
OUTPUT_DIR = r"c:\Projetos\Cosmo Alma TV\portal-egregora\backend\generated_contracts"

def generate_contract_docx(
    nome_completo: str, 
    nome_comercial: str, 
    razao_social: str, 
    cnpj_cpf: str, 
    email: str, 
    chave_pix: str, 
    youtube_id: str,
    genero: str = None,
    estado_civil: str = None,
    cep: str = None,
    endereco: str = None,
    cidade: str = None,
    uf: str = None,
    pais: str = None
) -> str:
    """
    Generates a personalized contract based on the template by replacing the default values.
    Returns the path to the generated contract file.
    """
    if not os.path.exists(OUTPUT_DIR):
        os.makedirs(OUTPUT_DIR)

    doc = docx.Document(TEMPLATE_PATH)

    # Gender adjustments
    genero_val = (genero or "").upper()
    nacionalidade = "brasileiro(a)"
    if genero_val == "M":
        nacionalidade = "brasileiro"
    elif genero_val == "F":
        nacionalidade = "brasileira"

    # Address block
    address_str = f"{endereco or ''}, {cidade or ''} - {uf or ''}, CEP {cep or ''}, {pais or 'Brasil'}"

    # Determine tags
    clean_doc = "".join([c for c in cnpj_cpf if c.isdigit()])
    is_cnpj = len(clean_doc) > 11
    tipo_pessoa = "jurídica" if is_cnpj else "física"
    cpf_ou_cnpj = "CNPJ" if is_cnpj else "CPF"
    nome_parte = razao_social or nome_completo

    # Replacements dictionary based on V3 bracketed placeholders
    replacements = {
        "[nome PESSOA FISICA OU JURÍDICA]": nome_parte.upper(),
        "[física/jurídica]": tipo_pessoa,
        "[CPF/CNPJ]": cpf_ou_cnpj,
        "[número cpf OU cnpj]": cnpj_cpf,
        "[NOME]": nome_completo.upper(),
        "[ESTADO civil]": (estado_civil or "solteiro(a)").lower(),
        "[CPF]": cnpj_cpf,
        "[ENDEREÇO]": address_str,
    }

    # Helper function to replace text in runs while preserving style
    def replace_text_in_paragraphs(paragraphs):
        for paragraph in paragraphs:
            for key, value in replacements.items():
                if key in paragraph.text:
                    # Replace in the whole paragraph text and reset runs to avoid breaking split text
                    text = paragraph.text.replace(key, value)
                    paragraph.text = text

    # Replace in main body paragraphs
    replace_text_in_paragraphs(doc.paragraphs)

    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_text_in_paragraphs(cell.paragraphs)

    # Save to generated contracts
    clean_cnpj_cpf = cnpj_cpf.replace('/', '').replace('.', '').replace('-', '')
    filename = f"contrato_{nome_comercial.lower().replace(' ', '_')}_{clean_cnpj_cpf}.docx"
    output_path = os.path.join(OUTPUT_DIR, filename)
    doc.save(output_path)
    return output_path
